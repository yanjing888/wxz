#!/usr/bin/env python3
"""Generate sample DOCX matching current ReportService output (v2)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name("实验总结报告-示例-牛顿环.docx")


def heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val
    return t


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.54)
        s.left_margin = s.right_margin = Cm(3.17)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("实验总结报告")
    r.bold = True
    r.font.size = Pt(18)

    table(doc, ["项目", "内容"], [
        ["实验名称", "牛顿环实验"],
        ["学生姓名", "张三"],
        ["班级", "物理 2301"],
        ["生成时间", "2026年06月15日 11:48:54"],
        ["会话编号", "1024"],
    ])

    doc.add_paragraph()
    heading(doc, "1. 实验步骤回顾")
    for i, (title, desc) in enumerate([
        ("仪器检查与光路调节", "检查光源、显微镜与支架，调节光路使环纹清晰。"),
        ("零级环心定位与视场居中", "定位干涉中心，记录鼓轮初读数。"),
        ("暗环直径测量", "读取各暗环左右读数并计算直径。"),
        ("曲率半径或波长计算", "代入公式计算 R 并比较误差。"),
    ], start=1):
        body(doc, f"{i}. {title}：{desc}")

    heading(doc, "2. 实操回顾与数据统计")
    table(doc, ["指标", "数值", "指标", "数值"], [
        ["有效纠错次数", "3", "纠错标注点数", "5"],
        ["教程查阅次数", "2", "严重告警 (L3)", "0"],
    ])

    heading(doc, "3. 操作纠错记录")
    table(doc, ["序号", "步骤", "错误类型", "时间", "问题描述"], [
        ["1", "仪器检查与光路调节", "光路未对准", "2026-06-11 09:15:22", "反射镜倾角过大，环纹偏心。"],
    ])

    heading(doc, "4. 实验数据记录")
    table(doc, ["步骤", "提交时间", "采集数据", "校验结果"], [
        ["暗环直径测量", "2026-06-11 09:45:00", "ring_m=10，reading_left_mm=12.345", "通过"],
    ])

    heading(doc, "5. 环境安全巡检记录")
    table(doc, ["巡检时间", "安全等级", "判断依据", "处置建议", "抽帧画面"], [
        ["2026-06-11 09:10:05", "L0 正常", "实验台整洁。", "—", "已采集"],
    ])

    heading(doc, "6. 知识巩固建议")
    body(doc, "• 等厚干涉与牛顿环明暗纹形成条件")
    body(doc, "• 环半径、直径与曲率半径、波长的关系式")

    heading(doc, "7. 后续学习路径")
    body(doc, "1. 复习教材「大学物理·光学」薄膜干涉与牛顿环")

    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("本报告由物小智智能实验平台根据本次实验过程数据自动生成。")
    fr.font.size = Pt(10)
    fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
